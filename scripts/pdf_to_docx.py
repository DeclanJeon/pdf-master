#!/usr/bin/env python3
"""PDF → DOCX 변환 v12: PyMuPDF + python-docx

핵심:
1. 모든 이미지 anchor behindDoc=1 → 텍스트가 항상 위(편집 가능)
2. 각 텍스트 라인별 폰트/사이즈 정확 보존
3. 표: Y/X 경계로 행높이/열너비, clip으로 셀 내용 추출
"""
import argparse, io, os, sys, logging
from pdf2docx import Converter

try:
    import fitz
except ImportError:
    print("PyMuPDF 미설치", file=sys.stderr); sys.exit(2)

try:
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("python-docx 미설치", file=sys.stderr); sys.exit(2)

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')
log = logging.getLogger(__name__)

PT_TO_MM = 25.4 / 72
MM_TO_PT = 72 / 25.4
DEFAULT_MARGIN_MM = 2
ANCHOR_PARA_HEIGHT = 4  # anchor용 빈 문단이 차지하는 대략적 높이(pt)


def _normalize_font(raw_font: str) -> str:
    lower = raw_font.lower()
    if "gungsuh" in lower or "궁서" in lower: return "Gungsuh"
    if "gulim" in lower or "굴림" in lower: return "Gulim"
    if "batang" in lower or "바탕" in lower: return "Batang"
    if "dotum" in lower or "돋움" in lower: return "Dotum"
    if "malgun" in lower or "맑은" in lower: return "Malgun Gothic"
    if "arial" in lower: return "Arial"
    if "times" in lower: return "Times New Roman"
    if "+" in raw_font:
        base = raw_font.split("+", 1)[1]
        return _normalize_font(base)
    return raw_font


def _color_to_hex(color) -> str | None:
    """Normalize PyMuPDF span/drawing colors to #RRGGBB."""
    if color is None:
        return None
    if isinstance(color, int):
        return f"#{color & 0xFFFFFF:06X}"
    try:
        if len(color) >= 3:
            r, g, b = color[:3]
            if all(isinstance(c, float) or (isinstance(c, int) and c <= 1) for c in (r, g, b)):
                return f"#{round(max(0, min(1, float(r))) * 255):02X}{round(max(0, min(1, float(g))) * 255):02X}{round(max(0, min(1, float(b))) * 255):02X}"
            return f"#{int(r) & 255:02X}{int(g) & 255:02X}{int(b) & 255:02X}"
    except Exception:
        return None
    return None


def _extract_cell_lines(page, clip_rect):
    """셀 영역의 텍스트를 라인별로 추출 (폰트/사이즈 포함)."""
    lines = []
    try:
        for b in page.get_text("dict", clip=clip_rect)["blocks"]:
            if b["type"] != 0: continue
            for line in b["lines"]:
                line_spans = []
                for span in line["spans"]:
                    text = span["text"]
                    if not text.strip() and not text: continue
                    line_spans.append({
                        "text": text,
                        "font": _normalize_font(span["font"]),
                        "size": span["size"],
                        "bold": "bold" in span["font"].lower(),
                    })
                if line_spans:
                    # 같은 라인의 스팬들을 하나의 라인으로 합치기
                    combined = {"text": "", "font": None, "size": None, "bold": False}
                    # 가장 긴 텍스트의 폰트/사이즈 기준
                    dominant = max(line_spans, key=lambda s: len(s["text"]))
                    combined["font"] = dominant["font"]
                    combined["size"] = dominant["size"]
                    combined["bold"] = dominant["bold"]
                    combined["text"] = "".join(s["text"] for s in line_spans).strip()
                    if combined["text"]:
                        lines.append(combined)
    except Exception:
        pass
    return lines


def _detect_align(page, clip_rect, is_merged=False):
    """셀 텍스트 정렬 감지."""
    if is_merged: return "center"
    try:
        for b in page.get_text("dict", clip=clip_rect)["blocks"]:
            if b["type"] != 0: continue
            for line in b["lines"]:
                for span in line["spans"]:
                    if span["text"].strip():
                        cw = clip_rect.x1 - clip_rect.x0
                        if cw > 0:
                            rel = (span["origin"][0] - clip_rect.x0) / cw
                            if 0.3 < rel < 0.7: return "center"
                            if rel > 0.8: return "right"
    except: pass
    return "left"


def analyze_page(page, mode: str) -> dict:
    pw, ph = page.rect.width, page.rect.height

    tables_raw = page.find_tables()
    table_elements = []
    table_rects = []
    for tab in tables_raw.tables:
        bbox = tab.bbox
        table_rects.append(fitz.Rect(bbox))
        data = tab.extract()
        y_vals = sorted(set(round(c[1],1) for c in tab.cells) | set(round(c[3],1) for c in tab.cells))
        x_vals = sorted(set(round(c[0],1) for c in tab.cells) | set(round(c[2],1) for c in tab.cells))
        row_heights = [y_vals[i+1]-y_vals[i] for i in range(len(y_vals)-1)]
        col_widths = [x_vals[i+1]-x_vals[i] for i in range(len(x_vals)-1)]

        # 각 셀의 라인별 내용 추출
        cell_details = []
        merged_rows = set()
        for ri in range(tab.row_count):
            is_merged = ri < len(data) and tab.col_count >= 2 and data[ri][1] is None
            if is_merged: merged_rows.add(ri)
            row_details = []
            for ci in range(tab.col_count):
                y0 = y_vals[ri]; y1 = y_vals[ri+1] if ri+1 < len(y_vals) else y0+30
                # 병합 행은 전체 폭으로 clip
                if is_merged:
                    x0 = x_vals[0]; x1 = x_vals[-1]
                else:
                    x0 = x_vals[ci]; x1 = x_vals[ci+1] if ci+1 < len(x_vals) else x0+200
                clip = fitz.Rect(x0, y0, x1, y1)
                
                lines = _extract_cell_lines(page, clip)
                align = _detect_align(page, clip, is_merged)
                
                # fallback: extract() 데이터
                if not lines:
                    raw_text = (data[ri][ci] or "") if ri < len(data) and ci < len(data[ri]) else ""
                    if raw_text:
                        lines = [{"text": raw_text, "font": "Gungsuh", "size": 15, "bold": False}]
                
                row_details.append({"lines": lines, "align": align})
            cell_details.append(row_details)

        table_elements.append({
            "data": data, "rows": tab.row_count, "cols": tab.col_count,
            "row_heights": row_heights, "col_widths": col_widths,
            "cell_details": cell_details, "merged_rows": merged_rows,
            "bbox": bbox,
        })

    # 표 바깥 텍스트
    text_elements = []
    blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
    for b in blocks:
        if b["type"] != 0: continue
        for line in b["lines"]:
            for span in line["spans"]:
                text = span["text"].strip()
                if not text: continue
                r = fitz.Rect(span["bbox"])
                if any(tr.contains(r) or tr.intersects(r) for tr in table_rects): continue
                font = _normalize_font(span["font"])
                # Absolute layout always uses explicit x indent (left). Flow modes
                # may still guess center/right from origin position.
                align = "left"
                if mode != "absolute":
                    if 0.35 < span["origin"][0] / pw < 0.65:
                        align = "center"
                    elif span["origin"][0] / pw > 0.7:
                        align = "right"
                text_elements.append({
                    "text": text,
                    "x": span["origin"][0],
                    "y": span["origin"][1],
                    "size": span["size"],
                    "font": font,
                    "align": align,
                    "bold": "bold" in span["font"].lower() or bool(span.get("flags", 0) & 16),
                    "color": _color_to_hex(span.get("color")),
                })

    # 모든 이미지 — 텍스트/표와 겹치는 것만 behindDoc 대상
    image_elements = []
    for info in page.get_image_info(xrefs=True):
        xref = info.get("xref")
        if xref is None: continue
        ir = fitz.Rect(info["bbox"])
        w_pt = ir.x1-ir.x0; h_pt = ir.y1-ir.y0
        if w_pt < 1 or h_pt < 1: continue
        try:
            img_data = page.parent.extract_image(xref)
            if not img_data or not img_data.get("image"): continue
        except: continue
        # 텍스트/표와 겹치는지 자동 판별
        in_table = any(fitz.Rect(t["bbox"]).intersects(ir) for t in table_elements)
        overlaps_text = any(ir.intersects(fitz.Rect(t.get("x",0)-2, t["y"]-2, t.get("x",0)+200, t["y"]+2))
                           for t in text_elements if "y" in t)
        needs_behind = in_table or overlaps_text
        image_elements.append({"xref": xref, "bbox": info["bbox"], "w_pt": w_pt, "h_pt": h_pt,
            "data": img_data["image"], "ext": img_data.get("ext","png"), "needs_behind": needs_behind})
    # 채움 사각형 등 벡터 도형 → 비트맵 앵커로 보존 (pdf2docx/절대 레이아웃 모두에서 색 박스 유실 방지)
    shape_elements = []
    try:
        drawings = page.get_drawings()
    except Exception:
        drawings = []
    for di, drawing in enumerate(drawings):
        fill = drawing.get("fill")
        rect = drawing.get("rect")
        if not fill or rect is None:
            continue
        try:
            x0, y0, x1, y1 = float(rect.x0), float(rect.y0), float(rect.x1), float(rect.y1)
        except Exception:
            continue
        w_pt = x1 - x0
        h_pt = y1 - y0
        if w_pt < 2 or h_pt < 2:
            continue
        # 페이지 전체 흰 배경 같은 큰 사각형은 제외
        if w_pt > pw * 0.95 and h_pt > ph * 0.95:
            continue
        fill_hex = _color_to_hex(fill)
        if not fill_hex or fill_hex.upper() in {"#FFFFFF", "#FFFFFE"}:
            continue
        # 네이티브 DrawingML 사각형 메타 (PNG 래스터 없이 색/위치 보존)
        try:
            stroke = drawing.get("color") or drawing.get("stroke")
            stroke_hex = _color_to_hex(stroke)
            if stroke_hex and stroke_hex.upper() == fill_hex.upper():
                stroke_hex = None
            shape_elements.append({
                "bbox": (x0, y0, x1, y1),
                "w_pt": w_pt,
                "h_pt": h_pt,
                "needs_behind": True,
                "kind": "vector-fill",
                "fill": fill_hex,
                "stroke": stroke_hex,
                "stroke_width": float(drawing.get("width") or 0.75),
            })
        except Exception as exc:
            log.debug("shape vector meta skip %s: %s", di, exc)

    # 이미지가 없는 경우에도 도형 이미지는 포함
    image_elements = shape_elements + image_elements


    return {"width": pw, "height": ph,
        "tables": table_elements, "texts": text_elements, "images": image_elements,
        "shapes": shape_elements}


def _hex_to_srgb(hex_color: str) -> str:
    """#RRGGBB → DrawingML srgbClr val (RRGGBB)."""
    value = (hex_color or "").strip().lstrip("#")
    if len(value) != 6:
        return "000000"
    return value.upper()


def _add_anchored_vector_shape_to_para(para, shape_el, doc_pr_counter):
    """문단에 DrawingML 사각형(채움/선) 앵커를 직접 추가.

    PNG 래스터 대신 네이티브 도형을 써서 Word에서 색 박스 편집이 가능하고
    해상도 손실이 없다.
    """
    fill_hex = shape_el.get("fill") or "#000000"
    stroke_hex = shape_el.get("stroke")
    x_emu = int(shape_el["bbox"][0] * 12700)
    y_emu = int(shape_el["bbox"][1] * 12700)
    w_emu = max(1, int(shape_el["w_pt"] * 12700))
    h_emu = max(1, int(shape_el["h_pt"] * 12700))
    counter = doc_pr_counter[0]
    doc_pr_counter[0] += 1
    behind = "1" if shape_el.get("needs_behind", True) else "0"
    rh = "0" if behind == "1" else str(counter)
    fill_val = _hex_to_srgb(fill_hex)
    stroke_xml = "<a:ln><a:noFill/></a:ln>"
    if stroke_hex:
        stroke_val = _hex_to_srgb(stroke_hex)
        if stroke_val != fill_val:
            stroke_w = max(9525, int(float(shape_el.get("stroke_width") or 0.75) * 12700))
            stroke_xml = (
                f'<a:ln w="{stroke_w}">'
                f'<a:solidFill><a:srgbClr val="{stroke_val}"/></a:solidFill>'
                f"</a:ln>"
            )

    # wps namespace for WordprocessingShape (DrawingML shape in Word).
    # python-docx nsmap does not register "wps", so declare URIs explicitly.
    drawing_xml = (
        '<w:drawing '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" '
        f'simplePos="0" relativeHeight="{rh}" behindDoc="{behind}" '
        f'locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{w_emu}" cy="{h_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f"<wp:wrapNone/>"
        f'<wp:docPr id="{counter}" name="Shape {counter}"/>'
        f"<wp:cNvGraphicFramePr/>"
        f"<a:graphic>"
        f'<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f"<wps:wsp>"
        f"<wps:cNvCnPr/>"
        f"<wps:cNvSpPr/>"
        f"<wps:spPr>"
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:solidFill><a:srgbClr val="{fill_val}"/></a:solidFill>'
        f"{stroke_xml}"
        f"</wps:spPr>"
        f"<wps:bodyPr/>"
        f"</wps:wsp>"
        f"</a:graphicData>"
        f"</a:graphic>"
        f"</wp:anchor>"
        f"</w:drawing>"
    )
    run = para.add_run()
    run._element.append(parse_xml(drawing_xml))


def _add_anchored_image_to_para(doc, para, img_el, doc_pr_counter):
    """문단에 inline 이미지를 추가한 후 anchor로 변환.
    텍스트/표와 겹치는 이미지만 behindDoc=1, 나머지는 전경(behindDoc=0).
    vector-fill 메타는 DrawingML 사각형으로 라우팅한다."""
    if img_el.get("kind") == "vector-fill" or not img_el.get("data"):
        if img_el.get("kind") == "vector-fill" or img_el.get("fill"):
            _add_anchored_vector_shape_to_para(para, img_el, doc_pr_counter)
        return

    behind = "1" if img_el.get("needs_behind") else "0"
    x_emu = int(img_el["bbox"][0] * 12700)
    y_emu = int(img_el["bbox"][1] * 12700)
    w_emu = int(img_el["w_pt"] * 12700)
    h_emu = int(img_el["h_pt"] * 12700)
    counter = doc_pr_counter[0]; doc_pr_counter[0] += 1

    run = para.add_run()
    run.add_picture(io.BytesIO(img_el["data"]),
                    width=Mm(img_el["w_pt"] * PT_TO_MM),
                    height=Mm(img_el["h_pt"] * PT_TO_MM))

    drawing = run._element.findall(qn('w:drawing'))[0]
    inline = drawing.find(qn('wp:inline'))
    if inline is None:
        return

    graphic = inline.find(qn('a:graphic'))
    graphicData = graphic.find(qn('a:graphicData'))
    pic = graphicData.find(qn('pic:pic'))
    rId = pic.find(qn('pic:blipFill')).find(qn('a:blip')).get(qn('r:embed'))

    # behindDoc=1 → relativeHeight=0 (최하위), behindDoc=0 → relativeHeight=counter (전경)
    rh = "0" if behind == "1" else str(counter)
    anchor_xml = (
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" '
        f'simplePos="0" relativeHeight="{rh}" behindDoc="{behind}" '
        f'locked="0" layoutInCell="1" allowOverlap="1" '
        f'{nsdecls("wp","r","a","pic")}>'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{w_emu}" cy="{h_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="{counter}" name="Picture {counter}"/>'
        f'<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic><pic:nvPicPr><pic:cNvPr id="{counter}" name="Picture {counter}"/>'
        f'<pic:cNvPicPr/><pic:nvGraphicFramePr><pic:graphicFrameLocks noChangeAspect="1"/></pic:nvGraphicFramePr></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rId}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
        f'</a:graphicData></a:graphic></wp:anchor>'
    )
    anchor_el = parse_xml(anchor_xml)
    drawing.replace(inline, anchor_el)


def _configure_section(section, page_data: dict, mode: str = "editable"):
    pw_pt = page_data["width"]; ph_pt = page_data["height"]
    # Absolute layout maps PDF page points directly; zero margins avoid a
    # horizontal/vertical shift relative to origin.
    margin_mm = 0.0 if mode == "absolute" else DEFAULT_MARGIN_MM
    section.page_width = Mm(pw_pt * PT_TO_MM)
    section.page_height = Mm(ph_pt * PT_TO_MM)
    section.top_margin = Mm(margin_mm)
    section.bottom_margin = Mm(margin_mm)
    section.left_margin = Mm(margin_mm)
    section.right_margin = Mm(margin_mm)



def _place_text_paragraph(doc: Document, tel: dict, cursor_y: float, page_data: dict, layout_mode: str, margin_pt: float):
    """Place one text line. Absolute mode sets left indent from PDF x."""
    txt_y = tel["y"]
    gap = max(0, txt_y - cursor_y)
    p = doc.add_paragraph()
    if gap > 1:
        p.paragraph_format.space_before = Pt(gap)
    run = p.add_run(tel["text"])
    _apply_font(run, tel)
    if layout_mode == "absolute":
        # Horizontal fidelity: indent from page left edge (margin is 0).
        x_pt = max(0.0, float(tel.get("x") or 0.0) - margin_pt)
        p.paragraph_format.left_indent = Pt(x_pt)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        if tel.get("align") == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif tel.get("align") == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(tel["size"] * 1.15)
    if any(
        img_el.get("needs_behind")
        and fitz.Rect(img_el["bbox"]).intersects(
            fitz.Rect(tel["x"] - 2, tel["y"] - 2, tel["x"] + 200, tel["y"] + tel["size"])
        )
        for img_el in page_data["images"]
    ):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="FFFFFF"/>')
        p._element.get_or_add_pPr().append(shd)
    return txt_y + tel["size"]


def _append_page(doc: Document, page_data: dict, mode: str, is_first_page: bool, doc_pr_counter: list):
    section = doc.sections[0] if is_first_page else doc.add_section(WD_SECTION.NEW_PAGE)
    # create_layout_docx maps absolute→editable for output_mode; recover intent
    # from page_data flag when present.
    layout_mode = page_data.get("layout_mode") or mode
    _configure_section(section, page_data, layout_mode)

    margin_pt = 0.0 if layout_mode == "absolute" else DEFAULT_MARGIN_MM * MM_TO_PT

    table_top = page_data["tables"][0]["bbox"][1] if page_data["tables"] else page_data["height"]

    # ── 1. 빈 문단에 모든 anchor 이미지 삽입 (behindDoc=1) ──
    anchor_para = doc.add_paragraph()
    anchor_para.paragraph_format.space_before = Pt(0)
    anchor_para.paragraph_format.space_after = Pt(0)
    anchor_para.paragraph_format.line_spacing = Pt(0.1)

    for img_el in page_data["images"]:
        _add_anchored_image_to_para(doc, anchor_para, img_el, doc_pr_counter)

    # ── 2. 표 위 텍스트 ──
    pre_texts = [t for t in page_data["texts"] if t["y"] < table_top]
    pre_texts.sort(key=lambda t: t["y"])

    cursor_y = margin_pt + ANCHOR_PARA_HEIGHT
    for tel in pre_texts:
        cursor_y = _place_text_paragraph(doc, tel, cursor_y, page_data, layout_mode, margin_pt)

    # ── 3. 표 ──
    for tinfo in page_data["tables"]:
        tab_y = tinfo["bbox"][1]
        gap = max(0, tab_y - cursor_y)
        if gap > 1:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(gap)
            spacer.paragraph_format.space_after = Pt(0)
            spacer.paragraph_format.line_spacing = Pt(0.1)
        table = _add_table(doc, tinfo)
        if layout_mode == "absolute" and table is not None:
            # Keep table left edge at PDF bbox x (page margin is 0).
            try:
                tab_x = max(0.0, float(tinfo["bbox"][0]) - margin_pt)
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                # Table left indent via tblInd
                tbl = table._tbl
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
                    tbl.insert(0, tblPr)
                existing = tblPr.find(qn('w:tblInd'))
                if existing is not None:
                    tblPr.remove(existing)
                # twips: 1pt = 20 twips
                tblPr.append(parse_xml(
                    f'<w:tblInd {nsdecls("w")} w:w="{int(tab_x * 20)}" w:type="dxa"/>'
                ))
            except Exception:
                pass
        cursor_y = tinfo["bbox"][3]

    # ── 4. 표 아래 텍스트 ──
    post_texts = [t for t in page_data["texts"] if t["y"] >= table_top]
    post_texts.sort(key=lambda t: t["y"])

    for tel in post_texts:
        cursor_y = _place_text_paragraph(doc, tel, cursor_y, page_data, layout_mode, margin_pt)


def build_docx(page_data: dict, mode: str) -> bytes:
    doc = Document()
    _append_page(doc, page_data, mode, True, [1])

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def create_absolute_layout_docx(pdf_doc, page_index: int = 0) -> bytes:
    page = pdf_doc[page_index]
    page_data = analyze_page(page, "absolute")
    page_data["layout_mode"] = "absolute"
    return build_docx(page_data, "absolute")


def create_layout_docx(pdf_doc, mode: str, start: int, end: int) -> bytes:
    page_count = len(pdf_doc)
    if page_count == 0:
        raise ValueError("PDF에 페이지가 없습니다")
    first = max(0, start)
    last = page_count - 1 if end < 0 else min(end, page_count - 1)
    if first > last:
        raise ValueError("변환할 PDF 페이지 범위가 올바르지 않습니다")

    doc = Document()
    doc_pr_counter = [1]
    for page_index in range(first, last + 1):
        page_mode = "absolute" if mode == "absolute" else mode
        # Keep absolute through placement so horizontal indents apply. Other
        # modes still use their flow builders.
        output_mode = page_mode
        page_data = analyze_page(pdf_doc[page_index], page_mode)
        page_data["layout_mode"] = page_mode
        _append_page(doc, page_data, output_mode, page_index == first, doc_pr_counter)

    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def create_faithful_docx_with_pdf2docx(input_path: str, output_path: str, start: int, end: int):
    cv = Converter(input_path)
    try:
        cv.convert(
            output_path,
            start=start,
            end=None if end < 0 else end,
            page_margin_factor_top=0.0,
            page_margin_factor_bottom=0.0,
            float_image_ignorable_gap=0.0,
            clip_image_res_ratio=8.0,
            extract_stream_table=True,
            shape_min_dimension=0.5,
            min_svg_w=0.5,
            min_svg_h=0.5,
            connected_border_tolerance=1.0,
            min_border_clearance=0.5,
            line_overlap_threshold=0.95,
            parse_lattice_table=True,
            parse_stream_table=True,
        )
    finally:
        cv.close()
    _inject_vector_fill_shapes_into_docx(input_path, output_path, start, end)

def _inject_vector_fill_shapes_into_docx(pdf_path: str, docx_path: str, start: int = 0, end: int = -1):
    """pdf2docx DOCX에 누락된 채움 사각형을 DrawingML 앵커 도형으로 보강."""
    try:
        pdf_doc = fitz.open(pdf_path)
    except Exception as exc:
        log.warning("vector fill inject: PDF open failed: %s", exc)
        return

    page_count = len(pdf_doc)
    first = max(0, start)
    last = page_count - 1 if end < 0 else min(end, page_count - 1)
    if first > last:
        pdf_doc.close()
        return

    page = pdf_doc[first]
    page_data = analyze_page(page, "absolute")
    shapes = page_data.get("shapes") or [
        img for img in page_data.get("images", [])
        if img.get("kind") == "vector-fill" or img.get("fill")
    ]
    pdf_doc.close()
    if not shapes:
        return

    try:
        doc = Document(docx_path)
    except Exception as exc:
        log.warning("vector fill inject: DOCX open failed: %s", exc)
        return

    body = doc.element.body
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = Pt(0.1)
    # move newly added paragraph to the top of the body
    body.insert(0, para._element)

    counter = [1000]
    for shape in shapes:
        try:
            _add_anchored_vector_shape_to_para(para, shape, counter)
        except Exception as exc:
            log.debug("vector fill inject skip: %s", exc)

    doc.save(docx_path)
    log.info("vector fill DrawingML shapes injected: %d", len(shapes))


def _add_table(doc: Document, tinfo: dict):
    rows = tinfo["rows"]; cols = tinfo["cols"]
    row_heights = tinfo["row_heights"]; col_widths = tinfo["col_widths"]
    cell_details = tinfo["cell_details"]; merged_rows = tinfo["merged_rows"]

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri in range(rows):
        if ri < len(row_heights) and row_heights[ri] > 0:
            table.rows[ri].height = Pt(row_heights[ri])
            table.rows[ri].height_rule = 1  # EXACTLY

    total_w = sum(col_widths)
    if total_w > 0:
        for ci in range(cols):
            if ci < len(col_widths):
                table.columns[ci].width = Mm(col_widths[ci] * PT_TO_MM)

    for ri in merged_rows:
        table.cell(ri, 0).merge(table.cell(ri, cols - 1))

    for ri in range(rows):
        for ci in range(cols):
            if ri in merged_rows and ci > 0: continue
            detail = cell_details[ri][ci] if ri < len(cell_details) and ci < len(cell_details[ri]) else None
            doc_cell = table.cell(ri, ci)
            # 기존 텍스트 제거
            for pp in doc_cell.paragraphs:
                for rr in pp.runs: rr.text = ""
            doc_cell.paragraphs[0].clear()

            if not detail or not detail.get("lines"):
                continue

            lines = detail["lines"]
            align = detail.get("align", "left")

            for li, line_info in enumerate(lines):
                if li > 0:
                    p = doc_cell.add_paragraph()
                else:
                    p = doc_cell.paragraphs[0]
                run = p.add_run(line_info["text"])
                _apply_font(run, line_info)
                if align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # 라인 간격: 폰트 사이즈 기반
                sz = line_info.get("size", 11)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = Pt(sz * 1.2)

            _set_cell_margins(doc_cell, top=10, bottom=10, start=30, end=30)


    return table


def _apply_font(run, fi: dict):
    sz = fi.get("size", 11); fn = fi.get("font", "Malgun Gothic")
    run.font.size = Pt(sz); run.font.name = fn
    if fi.get("bold"): run.bold = True
    color = fi.get("color")
    if isinstance(color, str) and color.startswith("#") and len(color) == 7:
        try:
            run.font.color.rgb = RGBColor(int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))
        except Exception:
            pass
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{fn}" w:hAnsi="{fn}" w:eastAsia="{fn}" w:cs="{fn}"/>')
        rPr.insert(0, rFonts)
    else:
        for attr in [qn('w:eastAsia'), qn('w:ascii'), qn('w:hAnsi'), qn('w:cs')]:
            rFonts.set(attr, fn)


def _set_cell_margins(cell, top=10, bottom=10, start=30, end=30):
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None: tcPr.remove(existing)
    tcPr.append(parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:start w:w="{start}" w:type="dxa"/>'
        f'<w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'))


def main():
    parser = argparse.ArgumentParser(description="PDF → DOCX 변환")
    parser.add_argument("input"); parser.add_argument("output")
    parser.add_argument("--layout-mode", choices=["faithful", "editable", "absolute"], default=None)
    parser.add_argument("--mode", choices=["faithful", "editable", "absolute"], default=None)
    parser.add_argument("--start", type=int, default=0)
    parser.add_argument("--end", type=int, default=-1)
    parser.add_argument("--margin", type=float, default=DEFAULT_MARGIN_MM)
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"입력 PDF를 찾을 수 없습니다: {args.input}", file=sys.stderr); sys.exit(3)

    try:
        pdf_doc = fitz.open(args.input)
    except Exception as e:
        print(f"PDF 열기 실패: {e}", file=sys.stderr); sys.exit(3)

    mode = args.layout_mode or args.mode or "faithful"
    log.info(f"변환 시작: {args.input} → {args.output} (모드={mode})")
    out_dir = os.path.dirname(os.path.abspath(args.output))
    os.makedirs(out_dir, exist_ok=True)

    if mode == "faithful":
        pdf_doc.close()
        create_faithful_docx_with_pdf2docx(args.input, args.output, args.start, args.end)
    else:
        docx_bytes = create_layout_docx(pdf_doc, mode, args.start, args.end)
        with open(args.output, "wb") as f:
            f.write(docx_bytes)
        pdf_doc.close()

    log.info(f"변환 완료: {args.output}")
    print(args.output)


if __name__ == "__main__":
    main()
